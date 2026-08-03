"""File processor — extracts text content from various file formats.

Supported formats:
- Excel (.xlsx) via openpyxl
- CSV (.csv) via csv module
- PDF (.pdf) via pdfplumber
- Word (.docx) via python-docx
- Images (.png, .jpg, .jpeg, .bmp, .tiff) via OCR (pytesseract + Pillow)
- JSON, XML, HTML — stdlib parsers
- Plain text (.txt, .md, .log, .py, .cs, .js, etc.)

Also provides file generation (Excel reports).
"""

from __future__ import annotations

import csv
import io
import json
import logging
import uuid
import defusedxml.ElementTree as ET  # Safe XML parsing (prevents XXE) — no fallback
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Optional

from pydantic import BaseModel

logger = logging.getLogger(__name__)

# Max text length to return to avoid blowing up LLM context
MAX_EXTRACTED_TEXT_LENGTH = 50_000
# Max rows to extract from tabular data
MAX_ROWS = 500


def _coerce_report_rows(data: Any) -> list[dict]:
    """Normalize LLM-supplied report data into a list of dict rows.

    Models don't reliably honor the ``[{col: val}, ...]`` row contract — they
    pass a list of strings, a list of lists, a columnar dict, or a bare scalar.
    Coerce every shape into ``list[dict]`` so Excel generation never crashes on
    ``data[0].keys()``. Well-formed dict rows pass through unchanged, so the
    happy path is byte-identical; this only turns a previous crash into output.
    """
    if data is None:
        return []
    # Columnar dict: {"Колонка": [v1, v2, ...]} → row-dicts (transpose)
    if isinstance(data, dict):
        if data and all(isinstance(v, (list, tuple)) for v in data.values()):
            cols = list(data.keys())
            n = max((len(v) for v in data.values()), default=0)
            return [{c: (data[c][i] if i < len(data[c]) else "") for c in cols}
                    for i in range(n)]
        return [data]  # single record
    if not isinstance(data, (list, tuple)):
        return [{"Значение": data}]  # bare scalar
    rows: list[dict] = []
    for item in data:
        if isinstance(item, dict):
            rows.append(item)
        elif isinstance(item, (list, tuple)):
            rows.append({f"Колонка {i + 1}": v for i, v in enumerate(item)})
        elif item is None:
            continue
        else:
            rows.append({"Значение": item})
    return rows


def _union_row_headers(rows: list[dict]) -> list[str]:
    """Order-preserving union of keys across ALL rows (not just ``rows[0]``), so
    heterogeneous rows don't silently drop columns present only in later rows."""
    headers: list[str] = []
    seen: set = set()
    for row in rows:
        for k in row.keys():
            if k not in seen:
                seen.add(k)
                headers.append(str(k))
    return headers

# Supported file extensions
TEXT_EXTENSIONS = {
    "txt", "md", "log", "py", "cs", "js", "ts", "jsx", "tsx",
    "css", "scss", "yaml", "yml", "toml", "ini", "cfg", "conf",
    "sh", "bat", "ps1", "sql", "r", "rb", "go", "rs", "java",
    "kt", "swift", "cpp", "c", "h", "hpp", "vb", "fs",
}

EXCEL_EXTENSIONS = {"xlsx", "xls"}
CSV_EXTENSIONS = {"csv", "tsv"}
PDF_EXTENSIONS = {"pdf"}
WORD_EXTENSIONS = {"docx"}
IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "bmp", "tiff", "tif", "gif", "webp"}
JSON_EXTENSIONS = {"json", "geojson"}
XML_EXTENSIONS = {"xml", "svg"}
HTML_EXTENSIONS = {"html", "htm"}

ALL_SUPPORTED = (
    TEXT_EXTENSIONS | EXCEL_EXTENSIONS | CSV_EXTENSIONS |
    PDF_EXTENSIONS | WORD_EXTENSIONS | IMAGE_EXTENSIONS |
    JSON_EXTENSIONS | XML_EXTENSIONS | HTML_EXTENSIONS
)


class FileExtractionResult(BaseModel):
    """Result of file content extraction."""
    success: bool
    text: str
    format: str  # "text", "table", "structured", "ocr", "image"
    rows_count: Optional[int] = None
    pages_count: Optional[int] = None
    error: Optional[str] = None
    # For images: base64 data URL ready for multimodal LLM input
    # (data:image/<mime>;base64,<...>). When present, callers should send
    # the image to the LLM directly rather than relying on `text`.
    image_data_url: Optional[str] = None
    image_filename: Optional[str] = None


class _HTMLTextExtractor(HTMLParser):
    """Extract visible text from HTML."""

    def __init__(self) -> None:
        super().__init__()
        self._text_parts: list[str] = []
        self._skip = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, Optional[str]]]) -> None:
        if tag in ("script", "style", "head"):
            self._skip = True

    def handle_endtag(self, tag: str) -> None:
        if tag in ("script", "style", "head"):
            self._skip = False
        if tag in ("p", "div", "br", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6", "td", "th"):
            self._text_parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._skip:
            self._text_parts.append(data.strip())

    def get_text(self) -> str:
        return " ".join(p for p in self._text_parts if p).strip()


class FileProcessor:
    """Extracts text from uploaded files and generates reports."""

    def __init__(self, files_dir: Optional[Path] = None):
        self._files_dir = files_dir

    @staticmethod
    def is_supported(filename: str) -> bool:
        """Check if a file extension is supported."""
        ext = _get_extension(filename)
        return ext in ALL_SUPPORTED

    def extract(self, data: bytes, filename: str) -> FileExtractionResult:
        """Extract text content from file data.

        Returns a FileExtractionResult with the extracted text.
        """
        MAX_EXTRACT_BYTES = 50 * 1024 * 1024  # 50MB absolute maximum
        if len(data) > MAX_EXTRACT_BYTES:
            return FileExtractionResult(
                success=False,
                text="",
                format="error",
                error="File too large for processing",
            )

        ext = _get_extension(filename)

        try:
            if ext in TEXT_EXTENSIONS:
                return self._extract_text(data, filename)
            elif ext in EXCEL_EXTENSIONS:
                return self._extract_excel(data, filename)
            elif ext in CSV_EXTENSIONS:
                return self._extract_csv(data, filename)
            elif ext in PDF_EXTENSIONS:
                return self._extract_pdf(data, filename)
            elif ext in WORD_EXTENSIONS:
                return self._extract_docx(data, filename)
            elif ext in IMAGE_EXTENSIONS:
                return self._extract_image(data, filename)
            elif ext in JSON_EXTENSIONS:
                return self._extract_json(data, filename)
            elif ext in XML_EXTENSIONS:
                return self._extract_xml(data, filename)
            elif ext in HTML_EXTENSIONS:
                return self._extract_html(data, filename)
            else:
                return FileExtractionResult(
                    success=False,
                    text="",
                    format="unknown",
                    error=f"Unsupported file format: .{ext}",
                )
        except Exception as e:
            logger.exception("File extraction error for %s", filename)
            return FileExtractionResult(
                success=False,
                text="",
                format="error",
                error=f"Failed to process file: {str(e)}",
            )

    def generate_excel(
        self,
        data: list[dict[str, Any]],
        filename: str = "report.xlsx",
        sheet_name: str = "Report",
    ) -> tuple[str, Path]:
        """Generate an Excel file from data rows.

        Returns (file_id, file_path).
        """
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        if not self._files_dir:
            raise RuntimeError("files_dir not configured")

        self._files_dir.mkdir(parents=True, exist_ok=True)

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name

        # Normalize model-supplied rows into list[dict] (defends against a bare
        # string/list/columnar-dict — otherwise data[0].keys() crashes).
        data = _coerce_report_rows(data)

        if not data:
            ws.append(["No data"])
            file_id = str(uuid.uuid4())[:8]
            path = self._files_dir / f"{file_id}_{filename}"
            wb.save(str(path))
            return file_id, path

        # Headers — union across all rows so later-row columns aren't dropped
        headers = _union_row_headers(data)
        ws.append(headers)

        # Style headers
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)
        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )

        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center")
            cell.border = thin_border

        # Data rows + detect numeric columns for formatting
        numeric_cols: set[int] = set()
        for row_data in data:
            row_values = [row_data.get(h, "") for h in headers]
            ws.append(row_values)
            # Detect which columns have numeric data
            for col_idx, val in enumerate(row_values, 1):
                if isinstance(val, (int, float)):
                    numeric_cols.add(col_idx)

        # Apply borders and number format to data cells
        for row_idx in range(2, len(data) + 2):
            for col_idx in range(1, len(headers) + 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.border = thin_border
                # Format numbers with thousand separator
                if col_idx in numeric_cols and isinstance(cell.value, (int, float)):
                    if isinstance(cell.value, float) and cell.value != int(cell.value):
                        cell.number_format = '#,##0.00'
                    else:
                        cell.number_format = '#,##0'
                    cell.alignment = Alignment(horizontal="right")

        # Add totals row for numeric columns (if there are any)
        if numeric_cols and len(data) >= 2:
            total_row = len(data) + 2
            total_font = Font(bold=True, size=11)
            total_fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")
            ws.cell(row=total_row, column=1, value="ИТОГО").font = total_font
            ws.cell(row=total_row, column=1).fill = total_fill
            ws.cell(row=total_row, column=1).border = thin_border
            for col_idx in range(2, len(headers) + 1):
                cell = ws.cell(row=total_row, column=col_idx)
                cell.border = thin_border
                cell.fill = total_fill
                if col_idx in numeric_cols:
                    # SUM formula
                    col_letter = openpyxl.utils.get_column_letter(col_idx)
                    cell.value = f"=SUM({col_letter}2:{col_letter}{total_row - 1})"
                    cell.font = total_font
                    cell.number_format = '#,##0.00'
                    cell.alignment = Alignment(horizontal="right")

        # Zebra striping for data rows
        stripe_fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
        for row_idx in range(2, len(data) + 2):
            if row_idx % 2 == 0:
                for col_idx in range(1, len(headers) + 1):
                    ws.cell(row=row_idx, column=col_idx).fill = stripe_fill

        # Auto-width columns
        for col_idx, header in enumerate(headers, 1):
            max_length = len(str(header))
            for row_idx in range(2, min(len(data) + 2, 102)):  # Check first 100 rows
                cell_value = ws.cell(row=row_idx, column=col_idx).value
                if cell_value:
                    max_length = max(max_length, len(str(cell_value)))
            ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = min(max_length + 3, 50)

        # Freeze header row
        ws.freeze_panes = "A2"

        # Auto-filter on header
        ws.auto_filter.ref = ws.dimensions

        file_id = str(uuid.uuid4())[:8]
        path = self._files_dir / f"{file_id}_{filename}"
        wb.save(str(path))
        logger.info("Generated Excel: %s (%d rows)", path, len(data))

        return file_id, path

    def generate_multi_sheet_excel(
        self,
        sheets: list[dict[str, Any]],
        filename: str = "report.xlsx",
    ) -> tuple[str, Path]:
        """Generate a multi-sheet Excel file.

        Each sheet: {"name": "Sheet Name", "data": [{"col": "val"}, ...]}
        Returns (file_id, file_path).
        """
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        if not self._files_dir:
            raise RuntimeError("files_dir not configured")

        self._files_dir.mkdir(parents=True, exist_ok=True)
        wb = openpyxl.Workbook()
        # Remove default sheet
        wb.remove(wb.active)

        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)
        stripe_fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
        total_fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")
        total_font = Font(bold=True, size=11)
        thin_border = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        )

        for sheet_info in sheets:
            sheet_name = sheet_info.get("name", "Sheet")[:31]  # Excel limit
            data = _coerce_report_rows(sheet_info.get("data", []))
            ws = wb.create_sheet(title=sheet_name)

            if not data:
                ws.append(["No data"])
                continue

            # Headers — union across all rows so later-row columns aren't dropped
            headers = _union_row_headers(data)
            ws.append(headers)
            for col_idx in range(1, len(headers) + 1):
                cell = ws.cell(row=1, column=col_idx)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center")
                cell.border = thin_border

            # Data rows
            numeric_cols: set[int] = set()
            for row_data in data:
                row_values = [row_data.get(h, "") for h in headers]
                ws.append(row_values)
                for col_idx, val in enumerate(row_values, 1):
                    if isinstance(val, (int, float)):
                        numeric_cols.add(col_idx)

            # Borders + number format + zebra
            for row_idx in range(2, len(data) + 2):
                for col_idx in range(1, len(headers) + 1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    cell.border = thin_border
                    if col_idx in numeric_cols and isinstance(cell.value, (int, float)):
                        cell.number_format = '#,##0.00' if isinstance(cell.value, float) and cell.value != int(cell.value) else '#,##0'
                        cell.alignment = Alignment(horizontal="right")
                if row_idx % 2 == 0:
                    for col_idx in range(1, len(headers) + 1):
                        ws.cell(row=row_idx, column=col_idx).fill = stripe_fill

            # Totals row for numeric columns
            if numeric_cols and len(data) >= 2:
                total_row = len(data) + 2
                ws.cell(row=total_row, column=1, value="ИТОГО").font = total_font
                ws.cell(row=total_row, column=1).fill = total_fill
                ws.cell(row=total_row, column=1).border = thin_border
                for col_idx in range(2, len(headers) + 1):
                    cell = ws.cell(row=total_row, column=col_idx)
                    cell.border = thin_border
                    cell.fill = total_fill
                    if col_idx in numeric_cols:
                        col_letter = openpyxl.utils.get_column_letter(col_idx)
                        cell.value = f"=SUM({col_letter}2:{col_letter}{total_row - 1})"
                        cell.font = total_font
                        cell.number_format = '#,##0.00'

            # Auto-width + freeze + filter
            for col_idx, header in enumerate(headers, 1):
                max_length = len(str(header))
                for row_idx in range(2, min(len(data) + 2, 52)):
                    val = ws.cell(row=row_idx, column=col_idx).value
                    if val:
                        max_length = max(max_length, len(str(val)))
                ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = min(max_length + 3, 50)
            ws.freeze_panes = "A2"
            ws.auto_filter.ref = ws.dimensions

        file_id = str(uuid.uuid4())[:8]
        path = self._files_dir / f"{file_id}_{filename}"
        wb.save(str(path))
        total_rows = sum(len(s.get("data", [])) for s in sheets)
        logger.info("Generated multi-sheet Excel: %s (%d sheets, %d rows)", path, len(sheets), total_rows)

        return file_id, path

    # --- Private extraction methods ---

    def _extract_text(self, data: bytes, filename: str) -> FileExtractionResult:
        """Extract plain text files."""
        text = _decode_text(data)
        text = _truncate(text)
        return FileExtractionResult(
            success=True,
            text=text,
            format="text",
        )

    def _extract_excel(self, data: bytes, filename: str) -> FileExtractionResult:
        """Extract data from Excel (.xlsx) files."""
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        total_rows = 0

        for sheet_name in wb.sheetnames:
            if total_rows >= MAX_ROWS:
                parts.append(f"... (remaining sheets skipped, {MAX_ROWS} rows shown)")
                break
            ws = wb[sheet_name]
            parts.append(f"=== Sheet: {sheet_name} ===")

            rows_in_sheet = 0
            for row in ws.iter_rows(values_only=True):
                if total_rows >= MAX_ROWS:
                    parts.append(f"... (truncated, showing {MAX_ROWS} rows)")
                    break
                # Convert each cell to string, handle None
                cells = [str(cell) if cell is not None else "" for cell in row]
                parts.append("\t".join(cells))
                rows_in_sheet += 1
                total_rows += 1

            if rows_in_sheet == 0:
                parts.append("(empty sheet)")

        wb.close()

        text = _truncate("\n".join(parts))
        return FileExtractionResult(
            success=True,
            text=text,
            format="table",
            rows_count=total_rows,
        )

    def _extract_csv(self, data: bytes, filename: str) -> FileExtractionResult:
        """Extract data from CSV/TSV files."""
        text_content = _decode_text(data)
        ext = _get_extension(filename)
        delimiter = "\t" if ext == "tsv" else ","

        reader = csv.reader(io.StringIO(text_content), delimiter=delimiter)
        parts: list[str] = []
        row_count = 0

        for row in reader:
            if row_count >= MAX_ROWS:
                parts.append(f"... (truncated, showing {MAX_ROWS} rows)")
                break
            parts.append("\t".join(row))
            row_count += 1

        text = _truncate("\n".join(parts))
        return FileExtractionResult(
            success=True,
            text=text,
            format="table",
            rows_count=row_count,
        )

    def _extract_pdf(self, data: bytes, filename: str) -> FileExtractionResult:
        """Extract text from PDF files."""
        import pdfplumber

        pdf = pdfplumber.open(io.BytesIO(data))
        parts: list[str] = []
        total_table_rows = 0
        page_count = len(pdf.pages)

        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                parts.append(f"--- Page {i + 1} ---")
                parts.append(text)

            # Also extract tables if present (cap at MAX_ROWS total)
            tables = page.extract_tables()
            for t_idx, table in enumerate(tables):
                if total_table_rows >= MAX_ROWS:
                    break
                parts.append(f"[Table {t_idx + 1} on page {i + 1}]")
                for row in table:
                    if total_table_rows >= MAX_ROWS:
                        parts.append(f"... (tables truncated at {MAX_ROWS} rows)")
                        break
                    cells = [str(cell) if cell is not None else "" for cell in row]
                    parts.append("\t".join(cells))
                    total_table_rows += 1

        pdf.close()

        text = _truncate("\n".join(parts))
        if not text.strip():
            text = f"[PDF: {filename}, {page_count} pages. No extractable text (might be scanned/image-based).]"

        return FileExtractionResult(
            success=True,
            text=text,
            format="text",
            pages_count=page_count,
        )

    def _extract_docx(self, data: bytes, filename: str) -> FileExtractionResult:
        """Extract text from Word (.docx) files."""
        import docx

        doc = docx.Document(io.BytesIO(data))
        parts: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Also extract tables
        for t_idx, table in enumerate(doc.tables):
            parts.append(f"\n[Table {t_idx + 1}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append("\t".join(cells))

        text = _truncate("\n".join(parts))
        return FileExtractionResult(
            success=True,
            text=text,
            format="text",
        )

    def _extract_image(self, data: bytes, filename: str) -> FileExtractionResult:
        """Prepare an uploaded image for multimodal LLM input.

        Returns a base64 data URL in `image_data_url` so the chat layer can
        attach it directly to the user message (Gemini-3-flash sees it
        natively — no OCR needed).
        """
        import base64

        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "png"
        # Map the file extension to a sensible image MIME type. Gemini accepts
        # png, jpeg, webp, gif natively. Fallback to png for unknown types
        # (still works since the receiver inspects bytes, not the MIME tag).
        mime_map = {
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg",
            "png": "image/png",
            "webp": "image/webp",
            "gif": "image/gif",
            "bmp": "image/bmp",
            "tiff": "image/tiff",
            "tif": "image/tiff",
        }
        mime = mime_map.get(ext, "image/png")
        b64 = base64.b64encode(data).decode("ascii")
        data_url = f"data:{mime};base64,{b64}"

        # Short text marker that ends up in chat history; the *real* image
        # bytes travel separately via image_data_url and only attach to the
        # current LLM call.
        marker = f"[Прикреплено изображение: {filename} ({len(data) // 1024} KB)]"

        return FileExtractionResult(
            success=True,
            text=marker,
            format="image",
            image_data_url=data_url,
            image_filename=filename,
        )

    def _extract_json(self, data: bytes, filename: str) -> FileExtractionResult:
        """Extract and pretty-print JSON."""
        text_content = _decode_text(data)
        parsed = json.loads(text_content)
        pretty = json.dumps(parsed, indent=2, ensure_ascii=False)
        pretty = _truncate(pretty)
        return FileExtractionResult(
            success=True,
            text=pretty,
            format="structured",
        )

    def _extract_xml(self, data: bytes, filename: str) -> FileExtractionResult:
        """Extract text content from XML."""
        text_content = _decode_text(data)
        root = ET.fromstring(text_content)
        parts: list[str] = []
        _xml_to_text(root, parts)
        text = _truncate("\n".join(parts))
        return FileExtractionResult(
            success=True,
            text=text,
            format="structured",
        )

    def _extract_html(self, data: bytes, filename: str) -> FileExtractionResult:
        """Extract visible text from HTML."""
        text_content = _decode_text(data)
        extractor = _HTMLTextExtractor()
        extractor.feed(text_content)
        text = _truncate(extractor.get_text())
        return FileExtractionResult(
            success=True,
            text=text,
            format="text",
        )


# --- Helpers ---

def _get_extension(filename: str) -> str:
    """Get lowercase file extension without the dot."""
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def _decode_text(data: bytes) -> str:
    """Decode bytes to string, trying multiple encodings."""
    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("latin-1")  # latin-1 always succeeds


def _truncate(text: str) -> str:
    """Truncate text to MAX_EXTRACTED_TEXT_LENGTH."""
    if len(text) > MAX_EXTRACTED_TEXT_LENGTH:
        return text[:MAX_EXTRACTED_TEXT_LENGTH] + "\n... [truncated]"
    return text


# --- Blocked file extensions (dangerous executables) ---
BLOCKED_EXTENSIONS = frozenset({
    "exe", "dll", "bat", "ps1", "cmd", "scr", "vbs",
    "com", "msi", "pif", "hta", "cpl", "inf", "reg",
})

# --- MIME magic bytes for file type validation ---
_MAGIC_BYTES: dict[str, list[tuple[bytes, int]]] = {
    "pdf": [(b"%PDF", 0)],
    "xlsx": [(b"PK\x03\x04", 0)],  # ZIP-based (Office Open XML)
    "docx": [(b"PK\x03\x04", 0)],
    "xls": [(b"\xd0\xcf\x11\xe0", 0)],  # OLE2
    "png": [(b"\x89PNG\r\n\x1a\n", 0)],
    "jpg": [(b"\xff\xd8\xff", 0)],
    "jpeg": [(b"\xff\xd8\xff", 0)],
    "gif": [(b"GIF87a", 0), (b"GIF89a", 0)],
    "bmp": [(b"BM", 0)],
    "tiff": [(b"II\x2a\x00", 0), (b"MM\x00\x2a", 0)],
    "tif": [(b"II\x2a\x00", 0), (b"MM\x00\x2a", 0)],
    "webp": [(b"RIFF", 0)],  # RIFF....WEBP
    "zip": [(b"PK\x03\x04", 0)],
    "xml": [(b"<?xml", 0), (b"\xef\xbb\xbf<?xml", 0)],  # with/without BOM
}


def validate_file_type(filename: str, data: bytes) -> tuple[bool, str]:
    """Validate file type by extension and magic bytes.

    Returns (is_valid, error_message).
    """
    ext = _get_extension(filename)

    # Block dangerous extensions
    if ext in BLOCKED_EXTENSIONS:
        return False, f"Тип файла .{ext} заблокирован по соображениям безопасности"

    # Universal check: block executables regardless of declared extension
    # Catches .exe renamed to .txt, etc.
    _EXECUTABLE_SIGNATURES = [
        (b"MZ", 0),              # PE/EXE/DLL (Windows)
        (b"\x7fELF", 0),         # ELF (Linux)
        (b"\xfe\xed\xfa", 0),    # Mach-O (macOS)
        (b"\xcf\xfa\xed\xfe", 0),  # Mach-O 64-bit
    ]
    for sig, offset in _EXECUTABLE_SIGNATURES:
        if len(data) >= offset + len(sig) and data[offset:offset + len(sig)] == sig:
            return False, f"Обнаружен исполняемый файл, замаскированный под .{ext}"

    # For known binary formats, verify magic bytes
    if ext in _MAGIC_BYTES:
        if not data:
            return False, f"Файл {filename} пуст"
        magic_list = _MAGIC_BYTES[ext]
        matched = False
        for magic, offset in magic_list:
            if len(data) >= offset + len(magic) and data[offset:offset + len(magic)] == magic:
                matched = True
                break
        if not matched:
            return False, f"Файл {filename} не соответствует формату .{ext} (проверка magic bytes)"

    return True, ""


class FileSandbox:
    """Managed upload directory with TTL-based cleanup."""

    def __init__(self, base_dir: str | Path, ttl_seconds: int = 3600):
        self.base_dir = Path(base_dir)
        self.ttl_seconds = ttl_seconds
        self.base_dir.mkdir(parents=True, exist_ok=True)

    def store(self, filename: str, content: bytes) -> str:
        """Store file and return file_id."""
        import time
        file_id = f"{int(time.time())}_{uuid.uuid4().hex[:8]}"
        # Sanitize filename: strip path components, keep only safe chars
        base_name = Path(filename).name  # strip directory components
        safe_name = "".join(
            c for c in base_name if c.isalnum() or c in ".-_"
        ) or "file"
        path = self.base_dir / f"{file_id}_{safe_name}"
        path.write_bytes(content)
        return file_id

    def get_path(self, file_id: str) -> Path | None:
        """Get file path by ID, None if expired or missing."""
        import time
        for path in self.base_dir.iterdir():
            if path.name.startswith(file_id):
                # Check TTL
                age = time.time() - path.stat().st_mtime
                if age > self.ttl_seconds:
                    try:
                        path.unlink()
                    except OSError:
                        pass
                    return None
                return path
        return None

    def cleanup_expired(self) -> int:
        """Remove files older than TTL. Returns count removed."""
        import time
        removed = 0
        cutoff = time.time() - self.ttl_seconds
        for path in self.base_dir.iterdir():
            if path.is_file() and path.stat().st_mtime < cutoff:
                try:
                    path.unlink()
                    removed += 1
                except OSError:
                    pass
        return removed


_XML_MAX_DEPTH = 100


def _xml_to_text(element: ET.Element, parts: list[str], depth: int = 0) -> None:
    """Recursively extract text from XML elements."""
    if depth >= _XML_MAX_DEPTH:
        parts.append(f"{'  ' * depth}... (max depth {_XML_MAX_DEPTH} reached)")
        return

    tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
    text = (element.text or "").strip()
    tail = (element.tail or "").strip()

    indent = "  " * depth
    if text:
        parts.append(f"{indent}<{tag}>: {text}")
    elif len(element) == 0:
        # Leaf element with attributes
        attrs = " ".join(f'{k}="{v}"' for k, v in element.attrib.items())
        if attrs:
            parts.append(f"{indent}<{tag} {attrs}>")
    else:
        parts.append(f"{indent}<{tag}>")

    for child in element:
        _xml_to_text(child, parts, depth + 1)

    if tail:
        parts.append(f"{indent}{tail}")
